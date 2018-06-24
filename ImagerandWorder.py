#/usr/bin/python3/
#coding=utf-8
#================ 简介 ===================
#     脚本：          伪·图像处理器     
#     作者：          北方重工NK1 
#     时间：          2018年1月22日 22:51:13
#     描述:           用于题图相关处理_作业帮
#================ 简介 ===================



from PIL import Image,ImageFilter,ImageEnhance
import datetime
from docx import Document
import os,sys,shutil
from docx.shared import Inches




class imorder:
    def __init__(self):
        pass

    def load(self,addressi):
        self.addressi=addressi
        self.image0=Image.open(self.addressi).filter(ImageFilter.DETAIL)#.filter(ImageFilter.DETAIL).filter(ImageFilter.DETAIL).filter(ImageFilter.DETAIL).filter(ImageFilter.DETAIL).filter(ImageFilter.DETAIL)
        self.image0=ImageEnhance.Sharpness (self.image0).enhance(3.0)#ImageEnhance.Contrast (self.image0).enhance(2)
        self.image0=ImageEnhance.Color(self.image0).enhance(0.1)
        #读取并将图像并锐化

    def cutbyelement(self,position,file):
        self.image0.crop(tuple([position[1],position[2],position[3],position[4]])).resize((400, position[4]-position[2]),Image.ANTIALIAS).save(file+'\screenshots\screenshot'+str(position[0])+'question.png')
        self.image0.crop(tuple([position[5],position[6],position[7],position[8]])).resize((400, position[8]-position[6]),Image.ANTIALIAS).save(file+'\screenshots\screenshot'+str(position[0])+'answer.png')
        #裁剪图片,剪成问题和答案两部分

'''
    用heightdown-heightup和width相减以取得图像大小元组size_0(width,height)
    研究word布局以确定最适大小元组standard(在窄页边距下,其width为分两栏能存放图片的width,其height为一列可以容纳图片高度的最大值),用除法与size_0比对值取整数为temp,
    
    temp只用做比对,在赋值时应按图片压缩处理        

    在width的比较中:
        若temp<=1则不变:
        若>=2则先压缩后等待height比较结果,若报错应报错于同一题号,
        报错时取得题号,时间,考点等与该题有关的所有信息置入文本文档log,并将log第一行计数器+1,返回到交互界面;
        否则直接压缩,不报错
    在height的比较中:
        若temp/height<=1则:
            若temp/(height/2)>=1:令图片的高=height/2
            否则若temp/(height/4)>=1:令图片的高=height/4
            否则若temp/(height/8)>=1:令图片的高=height/8
        若1.25>=tmep/height>=1:令图片的高=height
        若2>=temp/height>=1.25:将图片以height为界裁剪为两张图片并报错,方法同上
    
    待出现错误后补充本部分代码

'''
class FileSave:
    #文件读写模块
    def __init__(self,path,counter_in_render):
        
        if(os.path.exists(path+'/question'+str(str(datetime.datetime.now().year))+'.'+str(datetime.datetime.now().month)+'.docx')):
            pass
        else:
            shutil.copyfile(path+'/template.docx',path+'/question'+str(datetime.datetime.now().year)+'.'+str(datetime.datetime.now().month)+'.docx')
        if(os.path.exists(path+'/answer'+str(datetime.datetime.now().year)+'.'+str(datetime.datetime.now().month)+'.docx')):
            pass
        else:
            shutil.copyfile(path+'/template.docx',path+'/answer'+str(datetime.datetime.now().year)+'.'+str(datetime.datetime.now().month)+'.docx')

        print("FileSave has been launched.")
        self.path=path
        self.filename=path+'/result.txt'
        self.pointfile=path+'/point.txt'
        self.addressfile=path+'/address.txt'
        self.question=Document(path+'/question'+str(datetime.datetime.now().year)+'.'+str(datetime.datetime.now().month)+'.docx')
        print(path+'/question'+str(datetime.datetime.now().year)+'.'+str(datetime.datetime.now().month)+'.docx')
        self.answer=Document(path+'/answer'+str(datetime.datetime.now().year)+'.'+str(datetime.datetime.now().month)+'.docx')
        self.counter=list(range(counter_in_render))
        

    def write(self,string,address):
        string_0=str(string)
        address_0=str(address)
        point="考点="+string_0+";"+"网址="+address_0
        file=open(self.filename,'a')
        file.write(str(point))
        file.write('\n')
        file.close()
        file=open(self.pointfile,'a')
        file.write(string_0)
        file.write('\n')
        file.close()
        file=open(self.addressfile,'a')
        file.write(address_0)
        file.write('\n')
        file.close()

    def read(self):
        file=open(self.filename)
        for i in file.readlines():
            print(i)
        file.close()

    def writeword(self):
        paragraph=self.question.add_paragraph(str(datetime.datetime.now()))
        paragraph=self.answer.add_paragraph(str(datetime.datetime.now()))
        for i in self.counter:
            paragraph=self.question.add_paragraph(str(i+1))
            paragraph=self.answer.add_paragraph(str(i+1))
            self.question.add_picture(self.path+'\screenshots\screenshot'+str(i)+'question.png',width=Inches(3.95))
            self.answer.add_picture(self.path+'\screenshots\screenshot'+str(i)+'answer.png', width=Inches(3.95))
        self.question.save(self.path+'/question'+str(datetime.datetime.now().year)+'.'+str(datetime.datetime.now().month)+'.docx')
        self.answer.save(self.path+'/answer'+str(datetime.datetime.now().year)+'.'+str(datetime.datetime.now().month)+'.docx')


#以下为调试区:
'''
2018年2月4日 02:14:44
成功取得裁剪范围定点坐标
im=imorder()
im.load("H:/test.jpg",1,1)
im.cut()
'''
#file1=FileSave(sys.path[0],4)
